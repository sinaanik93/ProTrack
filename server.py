"""Small local server for the ProTrack MVP and real PDF/DOCX exports."""

from __future__ import annotations

import io
import hashlib
import hmac
import json
import mimetypes
import os
import re
import secrets
import time
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from http.cookies import SimpleCookie
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Flowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
LOGO = ROOT / "assets" / "protrack-official.png"
DATA_DIR = ROOT / "data"
AUTH_FILE = DATA_DIR / "auth.json"
HOST = "127.0.0.1"
PORT = int(os.environ.get("PROTRACK_PORT", "4173"))
SESSION_TTL = 8 * 60 * 60
SESSIONS: dict[str, dict] = {}


def password_hash(password: str, salt_hex: str) -> str:
    return hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), bytes.fromhex(salt_hex), 180_000).hex()


def default_auth_data() -> dict:
    defaults = [
        ("headcoach", "ProTrack2026!", "Head Coach", "head", None, None),
        ("assistant", "Coach2026!", "Sara Ahmadi", "assistant", "c2", None),
        ("player", "Player2026!", "Luca Marino", "player", None, "p1"),
    ]
    users = []
    for username, password, name, role, coach_id, player_id in defaults:
        salt = secrets.token_hex(16)
        users.append({"username": username, "name": name, "role": role, "coachId": coach_id, "playerId": player_id, "salt": salt, "hash": password_hash(password, salt), "active": True})
    return {"version": 1, "users": users}


def load_auth_data() -> dict:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    if not AUTH_FILE.exists():
        data = default_auth_data()
        AUTH_FILE.write_text(json.dumps(data, indent=2), encoding="utf-8")
        return data
    return json.loads(AUTH_FILE.read_text(encoding="utf-8"))


def save_auth_data(data: dict) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    AUTH_FILE.write_text(json.dumps(data, indent=2), encoding="utf-8")


def public_user(user: dict) -> dict:
    return {k: user.get(k) for k in ("username", "name", "role", "coachId", "playerId")}


def request_json(handler: SimpleHTTPRequestHandler, max_bytes: int = 100_000) -> dict:
    length = min(int(handler.headers.get("Content-Length", "0")), max_bytes)
    return json.loads(handler.rfile.read(length).decode("utf-8")) if length else {}


def session_user(handler: SimpleHTTPRequestHandler) -> dict | None:
    raw = handler.headers.get("Cookie", "")
    cookie = SimpleCookie(); cookie.load(raw)
    token = cookie.get("protrack_session")
    if not token: return None
    record = SESSIONS.get(token.value)
    if not record or record["expires"] < time.time():
        SESSIONS.pop(token.value, None); return None
    auth = load_auth_data()
    return next((u for u in auth["users"] if u["username"] == record["username"] and u.get("active", True)), None)


def clean_payload(handler: SimpleHTTPRequestHandler) -> dict:
    length = min(int(handler.headers.get("Content-Length", "0")), 4_000_000)
    raw = handler.rfile.read(length)
    data = json.loads(raw.decode("utf-8"))
    return {
        "title": str(data.get("title") or "ProTrack AI Analyst Report")[:180],
        "player": str(data.get("player") or "Player")[:100],
        "session": str(data.get("session") or "")[:30],
        "content": str(data.get("content") or "")[:500_000],
    }


def content_lines(content: str):
    for raw in content.replace("\r", "").split("\n"):
        line = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", raw).strip()
        if line:
            yield line


def is_heading(line: str) -> bool:
    known = {
        "Dashboard Database Output", "QA Validation Summary", "Rubric Engine Report",
        "Performance KPI Dashboard", "Padel-Specific Analysis", "Decision Making Analysis",
        "Error Analysis", "Match KPI Dashboard", "PDI Dashboard", "PTI Dashboard",
        "Assessment Dashboard", "Player Journey", "Progress Dashboard", "Results Generator",
        "Training Priority Matrix", "Development Plan", "Final Coach Report",
        "PROTRACK FINAL ASSESSMENT",
    }
    return line in known or bool(re.match(r"^\d{2}$", line))


class LogoCrop(Flowable):
    """Display the untouched official image through a clipped logo window."""

    def __init__(self, width=170 * mm, height=65 * mm):
        super().__init__()
        self.width = width
        self.height = height

    def draw(self):
        self.canv.saveState()
        clip = self.canv.beginPath()
        clip.rect(0, 0, self.width, self.height)
        self.canv.clipPath(clip, stroke=0, fill=0)
        full_height = self.width * (1670 / 942)
        crop_top = self.width * (680 / 942)
        draw_y = self.height + crop_top - full_height
        self.canv.drawImage(str(LOGO), 0, draw_y, width=self.width, height=full_height, preserveAspectRatio=True, mask="auto")
        self.canv.restoreState()


def build_docx(data: dict) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(.65)
    section.bottom_margin = Inches(.65)
    section.left_margin = Inches(.72)
    section.right_margin = Inches(.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(45, 55, 70)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        shape = title.add_run().add_picture(str(LOGO), width=Inches(6.4))
        shape.height = Inches(2.55)
        src_rect = OxmlElement("a:srcRect")
        src_rect.set("t", "40500")
        src_rect.set("b", "36500")
        shape._inline.graphic.graphicData.pic.blipFill.insert(1, src_rect)
    else:
        run = title.add_run("PROTRACK")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(15, 105, 235)
    subtitle = doc.add_paragraph("AI ANALYST · V2.0 FINAL MASTER CLEAN")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(8)
    subtitle.runs[0].font.color.rgb = RGBColor(105, 118, 136)
    heading = doc.add_heading(data["title"], level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Your Game. Elevated.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    pending_number = None
    for line in content_lines(data["content"]):
        if re.fullmatch(r"\d{2}", line):
            pending_number = line
            continue
        if is_heading(line):
            text = f"{pending_number}. {line}" if pending_number else line
            pending_number = None
            doc.add_heading(text, level=1)
        elif line.startswith("{") or line.startswith('"') or line in ("}", "},"):
            p = doc.add_paragraph(line)
            p.style = doc.styles["No Spacing"]
            p.runs[0].font.name = "Consolas"
            p.runs[0].font.size = Pt(7.5)
            p.runs[0].font.color.rgb = RGBColor(35, 82, 130)
        elif len(line) < 70 and (line.endswith(":") or " · " in line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.color.rgb = RGBColor(25, 71, 125)
        else:
            doc.add_paragraph(line)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PROTRACK · Your Game. Elevated. · v2.0 LOCKED").font.size = Pt(7)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_pdf(data: dict) -> bytes:
    output = io.BytesIO()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("PTTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=colors.HexColor("#0F69EB"), alignment=TA_CENTER, spaceAfter=10)
    sub_style = ParagraphStyle("PTSub", parent=styles["Normal"], fontSize=8, leading=12, textColor=colors.HexColor("#64748B"), alignment=TA_CENTER, spaceAfter=12)
    h_style = ParagraphStyle("PTH", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#0B376C"), spaceBefore=12, spaceAfter=7, keepWithNext=True)
    body_style = ParagraphStyle("PTBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=11.5, textColor=colors.HexColor("#263244"), spaceAfter=4)
    code_style = ParagraphStyle("PTCode", parent=body_style, fontName="Courier", fontSize=6.8, leading=9, leftIndent=4*mm, textColor=colors.HexColor("#285A91"))

    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=16*mm, leftMargin=16*mm, topMargin=16*mm, bottomMargin=17*mm,
                            title=data["title"], author="ProTrack AI Analyst v2.0")
    story = []
    if LOGO.exists():
        story.extend([LogoCrop(), Spacer(1, 7 * mm)])
    else:
        story.append(Paragraph("PROTRACK", title_style))
    story.extend([Paragraph("AI ANALYST · V2.0 FINAL MASTER CLEAN", sub_style), Spacer(1, 5*mm), Paragraph(data["title"], title_style), Paragraph("Your Game. Elevated.", sub_style), PageBreak()])
    pending_number = None
    for line in content_lines(data["content"]):
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if re.fullmatch(r"\d{2}", line):
            pending_number = line
            continue
        if is_heading(line):
            text = f"{pending_number}. {safe}" if pending_number else safe
            pending_number = None
            story.append(Paragraph(text, h_style))
        elif line.startswith("{") or line.startswith('"') or line in ("}", "},"):
            story.append(Paragraph(safe, code_style))
        else:
            story.append(Paragraph(safe, body_style))

    def page(canvas, document):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D5DCE6"))
        canvas.line(16*mm, 13*mm, 194*mm, 13*mm)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#718096"))
        canvas.drawString(16*mm, 8.5*mm, "PROTRACK · Your Game. Elevated. · v2.0 LOCKED")
        canvas.drawRightString(194*mm, 8.5*mm, f"Page {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page, onLaterPages=page)
    return output.getvalue()


class Handler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def log_message(self, fmt, *args):
        print(f"[ProTrack] {self.address_string()} {fmt % args}")

    def json_response(self, status: int, data: dict, cookie: str | None = None):
        payload = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        if cookie: self.send_header("Set-Cookie", cookie)
        self.send_header("Content-Length", str(len(payload)))
        self.end_headers(); self.wfile.write(payload)

    def do_GET(self):
        path = urlparse(self.path).path
        if path == "/api/auth/session":
            user = session_user(self)
            self.json_response(200, {"authenticated": True, "user": public_user(user)}) if user else self.json_response(401, {"authenticated": False})
            return
        super().do_GET()

    def do_POST(self):
        path = urlparse(self.path).path
        if path == "/api/auth/login":
            try:
                data = request_json(self); username = str(data.get("username", "")).strip().lower(); password = str(data.get("password", ""))
                auth = load_auth_data(); user = next((u for u in auth["users"] if u["username"] == username and u.get("active", True)), None)
                if not user or not hmac.compare_digest(user["hash"], password_hash(password, user["salt"])):
                    self.json_response(401, {"ok": False, "message": "Invalid credentials"}); return
                token = secrets.token_urlsafe(32); SESSIONS[token] = {"username": username, "expires": time.time() + SESSION_TTL}
                cookie = f"protrack_session={token}; HttpOnly; SameSite=Strict; Path=/; Max-Age={SESSION_TTL}"
                self.json_response(200, {"ok": True, "user": public_user(user)}, cookie); return
            except Exception:
                self.json_response(400, {"ok": False, "message": "Invalid request"}); return
        if path == "/api/auth/logout":
            raw = self.headers.get("Cookie", ""); cookie = SimpleCookie(); cookie.load(raw); item = cookie.get("protrack_session")
            if item: SESSIONS.pop(item.value, None)
            self.json_response(200, {"ok": True}, "protrack_session=; HttpOnly; SameSite=Strict; Path=/; Max-Age=0"); return
        if path == "/api/auth/password":
            user = session_user(self)
            if not user: self.json_response(401, {"ok": False}); return
            try:
                data = request_json(self); old = str(data.get("oldPassword", "")); new = str(data.get("newPassword", ""))
                if len(new) < 10 or not hmac.compare_digest(user["hash"], password_hash(old, user["salt"])):
                    self.json_response(400, {"ok": False, "message": "Password requirements not met"}); return
                auth = load_auth_data(); stored = next(u for u in auth["users"] if u["username"] == user["username"]); salt = secrets.token_hex(16); stored["salt"] = salt; stored["hash"] = password_hash(new, salt); save_auth_data(auth)
                self.json_response(200, {"ok": True}); return
            except Exception:
                self.json_response(400, {"ok": False}); return
        if path not in ("/api/export/pdf", "/api/export/docx"):
            self.send_error(404)
            return
        if not session_user(self):
            self.json_response(401, {"ok": False, "message": "Authentication required"})
            return
        try:
            data = clean_payload(self)
            if path.endswith("pdf"):
                payload, mime, extension = build_pdf(data), "application/pdf", "pdf"
            else:
                payload, mime, extension = build_docx(data), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
            filename = re.sub(r"[^A-Za-z0-9_-]+", "_", f"ProTrack_{data['player']}_Session_{data['session']}").strip("_")
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Disposition", f'attachment; filename="{filename}.{extension}"')
            self.send_header("Content-Length", str(len(payload)))
            self.end_headers()
            self.wfile.write(payload)
        except Exception as exc:
            self.send_error(400, explain=str(exc))


if __name__ == "__main__":
    mimetypes.add_type("text/javascript", ".js")
    mimetypes.add_type("application/manifest+json", ".webmanifest")
    mimetypes.add_type("application/manifest+json", ".json")
    mimetypes.add_type("image/svg+xml", ".svg")
    print(f"ProTrack AI Analyst is running at http://{HOST}:{PORT}")
    ThreadingHTTPServer((HOST, PORT), Handler).serve_forever()
